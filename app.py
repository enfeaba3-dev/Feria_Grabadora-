import atexit
import io
import json
import logging
import os
import re
import socket
import subprocess
import sys
import tempfile
import threading
import time
import uuid
import webbrowser
from datetime import datetime
from pathlib import Path
from typing import Any

from logging_setup import LOG_DIR, configure_logging, get_file_logger, tail_log

LOGGER = configure_logging("app")
APP_DIR = Path(__file__).resolve().parent
RUNTIME_DIR = APP_DIR / "runtime"
RUNTIME_DIR.mkdir(parents=True, exist_ok=True)
AGENT_STATE_PATH = RUNTIME_DIR / "agent_state.json"
HISTORY_DIR = APP_DIR / "transcripciones"
HISTORY_DIR.mkdir(parents=True, exist_ok=True)
HISTORY_MAX = 100  # máximo de entradas en el historial

# Idiomas soportados (Whisper)
LANGUAGE_NAMES = {
    "es": "Español",
    "en": "English",
    "fr": "Français",
    "de": "Deutsch",
    "it": "Italiano",
    "pt": "Português",
    "ca": "Català",
    "gl": "Galego",
    "eu": "Euskara",
    "zh": "中文",
    "ja": "日本語",
    "ru": "Русский",
    "ar": "العربية",
    "hi": "हिन्दी",
    "ko": "한국어",
    "nl": "Nederlands",
    "pl": "Polski",
    "tr": "Türkçe",
    "sv": "Svenska",
    "it": "Italiano",
    "uk": "Українська",
    "cs": "Čeština",
    "fi": "Suomi",
    "da": "Dansk",
    "no": "Norsk",
    "el": "Ελληνικά",
    "he": "עברית",
    "th": "ไทย",
    "vi": "Tiếng Việt",
    "id": "Bahasa Indonesia",
    "ms": "Bahasa Melayu",
    "ro": "Română",
    "hu": "Magyar",
    "sk": "Slovenčina",
    "bg": "Български",
}

try:
    from flask import Flask, g, jsonify, render_template, request, send_file
    from werkzeug.exceptions import RequestEntityTooLarge
    from werkzeug.utils import secure_filename
except Exception as exc:
    LOGGER.exception("Flask no está instalado")
    print("\nERROR: faltan dependencias. Ejecuta instalar.bat.\n")
    raise

from audio_pipeline import prepare_audio
from config_manager import (
    CONFIG_PATH,
    LANGUAGE_PATTERN,
    MODEL_IDS,
    load_config,
    save_config,
)
from diagnostics import create_support_bundle, run_diagnostics
from model_service import ModelService

AVAILABLE_MODELS = [
    {"id": "tiny", "name": "Tiny", "hint": "Velocidad máxima", "size": "~75 MB"},
    {"id": "base", "name": "Base", "hint": "Muy ligero", "size": "~145 MB"},
    {"id": "small", "name": "Small", "hint": "Buen equilibrio", "size": "~490 MB"},
    {"id": "medium", "name": "Medium", "hint": "Alta precisión", "size": "~1.5 GB"},
    {
        "id": "large-v2",
        "name": "Large v2",
        "hint": "Precisión avanzada",
        "size": "~3 GB",
    },
    {"id": "large-v3", "name": "Large v3", "hint": "Máxima calidad", "size": "~3 GB"},
    {
        "id": "large-v3-turbo",
        "name": "Large v3 Turbo",
        "hint": "Recomendado",
        "size": "~1.6 GB",
    },
]

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 1536 * 1024 * 1024
model_service = ModelService()


def error_response(
    code: str, message: str, status: int = 500, detail: str | None = None
):
    payload = {
        "error": {"code": code, "message": message},
        "request_id": getattr(g, "request_id", None),
    }
    if detail and app.debug:
        payload["error"]["detail"] = detail
    return jsonify(payload), status


class AgentController:
    def __init__(self) -> None:
        self.process: subprocess.Popen | None = None
        self.server_url = ""
        self.lock = threading.RLock()
        self.desired_running = False
        self.restart_times = []
        self.monitor_started = False

    def configure(self, server_url: str) -> None:
        self.server_url = server_url.rstrip("/")

    def start_monitor(self) -> None:
        with self.lock:
            if self.monitor_started:
                return
            self.monitor_started = True

        def monitor():
            while True:
                time.sleep(4.0)
                with self.lock:
                    if not self.desired_running:
                        continue
                    crashed = not self.process or self.process.poll() is not None
                    if not crashed:
                        continue
                    now = time.time()
                    self.restart_times = [
                        item for item in self.restart_times if now - item < 60
                    ]
                    if len(self.restart_times) >= 3:
                        LOGGER.error(
                            "Agente en bucle de fallos; reinicio automático suspendido"
                        )
                        self.desired_running = False
                        continue
                    self.restart_times.append(now)
                    LOGGER.warning(
                        "Agente detenido inesperadamente; reinicio automático %s/3",
                        len(self.restart_times),
                    )
                try:
                    self.start()
                except Exception:
                    LOGGER.exception("El reinicio automático del agente falló")

        threading.Thread(target=monitor, name="agent-monitor", daemon=True).start()

    def start(self) -> dict[str, Any]:
        with self.lock:
            config = load_config()
            if not config["dictation"]["enabled"]:
                self.desired_running = False
                return {
                    "status": "disabled",
                    "message": "El dictado global está desactivado.",
                }
            if os.name != "nt":
                self.desired_running = False
                return {
                    "status": "unsupported",
                    "message": "El dictado global está diseñado para Windows.",
                }
            self.desired_running = True
            if self.process and self.process.poll() is None:
                return self.status()
            command = [
                sys.executable,
                str(APP_DIR / "agent" / "dictation_agent.py"),
                "--server-url",
                self.server_url,
                "--state-path",
                str(AGENT_STATE_PATH),
                "--config-path",
                str(CONFIG_PATH),
            ]
            LOGGER.info("Iniciando agente | command=%s", command)
            creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
            agent_env = os.environ.copy()
            for d in [
                APP_DIR
                / ".venv"
                / "Lib"
                / "site-packages"
                / "nvidia"
                / "cublas"
                / "bin",
                APP_DIR
                / ".venv"
                / "Lib"
                / "site-packages"
                / "nvidia"
                / "cudnn"
                / "bin",
            ]:
                if d.exists():
                    agent_env["PATH"] = f"{d}{os.pathsep}{agent_env.get('PATH', '')}"
            self.process = subprocess.Popen(
                command,
                cwd=str(APP_DIR),
                env=agent_env,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                creationflags=creationflags,
            )
            return {"status": "starting", "pid": self.process.pid}

    def stop(self, keep_desired: bool = False) -> dict[str, Any]:
        with self.lock:
            if not keep_desired:
                self.desired_running = False
            if not self.process or self.process.poll() is not None:
                self.process = None
                return {"status": "stopped"}
            pid = self.process.pid
            LOGGER.info("Deteniendo agente | pid=%s", pid)
            self.process.terminate()
            try:
                self.process.wait(timeout=5)
            except subprocess.TimeoutExpired:
                LOGGER.warning("El agente no terminó; se fuerza cierre | pid=%s", pid)
                self.process.kill()
                self.process.wait(timeout=3)
            self.process = None
            return {"status": "stopped", "pid": pid}

    def restart(self) -> dict[str, Any]:
        self.stop(keep_desired=True)
        self.desired_running = True
        try:
            AGENT_STATE_PATH.unlink(missing_ok=True)
        except OSError:
            pass
        time.sleep(0.15)
        return self.start()

    def status(self) -> dict[str, Any]:
        with self.lock:
            process_running = bool(self.process and self.process.poll() is None)
            state = {}
            if AGENT_STATE_PATH.exists():
                try:
                    state = json.loads(AGENT_STATE_PATH.read_text(encoding="utf-8"))
                except Exception:
                    LOGGER.exception("No se pudo leer agent_state.json")
            state["process_running"] = process_running
            state["desired_running"] = self.desired_running
            state["automatic_restarts_last_minute"] = len(
                [item for item in self.restart_times if time.time() - item < 60]
            )
            state["supervisor_pid"] = self.process.pid if process_running else None
            updated = state.get("updated_at")
            state["heartbeat_age_seconds"] = (
                round(time.time() - updated, 2)
                if isinstance(updated, (int, float))
                else None
            )
            if (
                process_running
                and state.get("heartbeat_age_seconds")
                and state["heartbeat_age_seconds"] > 8
            ):
                state["status"] = "unresponsive"
            if not process_running and state.get("status") not in {
                "disabled",
                "unsupported",
            }:
                state["status"] = "stopped"
            return state


agent_controller = AgentController()
atexit.register(agent_controller.stop)


@app.before_request
def before_request():
    g.request_id = uuid.uuid4().hex[:10]
    g.request_started = time.perf_counter()
    LOGGER.info(
        "HTTP start | id=%s | %s %s | remote=%s | length=%s",
        g.request_id,
        request.method,
        request.path,
        request.remote_addr,
        request.content_length,
    )


@app.after_request
def after_request(response):
    elapsed = (
        round((time.perf_counter() - g.request_started) * 1000, 2)
        if hasattr(g, "request_started")
        else None
    )
    response.headers["X-Request-ID"] = getattr(g, "request_id", "")
    response.headers["Cache-Control"] = "no-store"
    LOGGER.info(
        "HTTP done | id=%s | %s %s | status=%s | ms=%s",
        getattr(g, "request_id", ""),
        request.method,
        request.path,
        response.status_code,
        elapsed,
    )
    return response


@app.get("/")
def index():
    config = load_config()
    return render_template("index.html", models=AVAILABLE_MODELS, config=config)


@app.get("/api/status")
def api_status():
    return jsonify(
        {
            "ready": True,
            "version": "2.2.0",
            "platform": sys.platform,
            "config": load_config(),
            "model": model_service.status(),
            "agent": agent_controller.status(),
            "request_id": g.request_id,
        }
    )


@app.get("/api/config")
def api_get_config():
    return jsonify({"config": load_config(), "request_id": g.request_id})


@app.put("/api/config")
def api_save_config():
    candidate = request.get_json(silent=True)
    if not isinstance(candidate, dict):
        return error_response(
            "CONFIG_INVALID", "La configuración enviada no es válida.", 400
        )
    try:
        config, warnings = save_config(candidate)
        result = (
            agent_controller.restart()
            if config["dictation"]["enabled"]
            else agent_controller.stop()
        )
        model_service.warmup_async(config["model"], config["device"])
        return jsonify(
            {
                "config": config,
                "warnings": warnings,
                "agent": result,
                "request_id": g.request_id,
            }
        )
    except Exception as exc:
        LOGGER.exception("No se pudo guardar la configuración")
        return error_response(
            "CONFIG_SAVE_FAILED", "No se pudo guardar la configuración.", 500, str(exc)
        )


@app.post("/api/agent/<action>")
def api_agent_action(action: str):
    if action == "start":
        result = agent_controller.start()
    elif action == "stop":
        result = agent_controller.stop()
    elif action == "restart":
        result = agent_controller.restart()
    else:
        return error_response(
            "AGENT_ACTION_INVALID", "Acción de agente no válida.", 404
        )
    return jsonify({"agent": result, "request_id": g.request_id})


@app.post("/api/model/warmup")
def api_model_warmup():
    data = request.get_json(silent=True) or {}
    config = load_config()
    model = data.get("model", config["model"])
    device = data.get("device", config["device"])
    if model not in MODEL_IDS:
        return error_response(
            "MODEL_INVALID", "El modelo seleccionado no es válido.", 400
        )
    return jsonify(
        {
            "warmup": model_service.warmup_async(model, device),
            "request_id": g.request_id,
        }
    )


@app.get("/api/audio-devices")
def api_audio_devices():
    try:
        import sounddevice as sd

        devices = []
        for index, item in enumerate(sd.query_devices()):
            if item.get("max_input_channels", 0) > 0:
                devices.append(
                    {
                        "id": index,
                        "name": item.get("name", f"Dispositivo {index}"),
                        "channels": item.get("max_input_channels", 0),
                        "default_samplerate": item.get("default_samplerate"),
                    }
                )
        return jsonify(
            {
                "devices": devices,
                "default": sd.default.device,
                "request_id": g.request_id,
            }
        )
    except Exception as exc:
        LOGGER.exception("No se pudieron consultar los dispositivos de audio")
        return error_response(
            "AUDIO_DEVICES_FAILED",
            "No se pudieron consultar los micrófonos.",
            500,
            str(exc),
        )


@app.post("/api/transcribe")
def transcribe():
    if "audio" not in request.files:
        return error_response("AUDIO_MISSING", "No se recibió ningún audio.", 400)
    audio_file = request.files["audio"]
    if not audio_file or not audio_file.filename:
        return error_response("AUDIO_EMPTY", "El archivo de audio está vacío.", 400)
    config = load_config()
    model_name = request.form.get("model", config["model"]).strip()
    if model_name not in MODEL_IDS:
        return error_response(
            "MODEL_INVALID", "El modelo seleccionado no es válido.", 400
        )
    language = request.form.get("language", config["language"]).strip().lower() or None
    if language and not LANGUAGE_PATTERN.fullmatch(language):
        return error_response(
            "LANGUAGE_INVALID", "El código de idioma no es válido.", 400
        )
    mode = request.form.get("mode", "file").strip().lower()
    if mode not in {"file", "live", "dictation"}:
        mode = "file"
    device = request.form.get("device", config["device"]).strip().lower()
    if device not in {"auto", "cpu", "cuda"}:
        device = "auto"
    safe_name = secure_filename(audio_file.filename) or "audio.bin"
    suffix = Path(safe_name).suffix.lower() or ".bin"
    started = time.perf_counter()
    LOGGER.info(
        "Audio recibido | id=%s | name=%s | mode=%s | model=%s | language=%s | device=%s",
        g.request_id,
        safe_name,
        mode,
        model_name,
        language or "auto",
        device,
    )
    try:
        with tempfile.TemporaryDirectory(prefix="feria_transcriber_") as temp_name:
            temp_dir = Path(temp_name)
            source = temp_dir / f"entrada{suffix}"
            audio_file.save(source)
            if not source.exists() or source.stat().st_size == 0:
                return error_response(
                    "AUDIO_EMPTY", "El audio recibido no contiene datos.", 400
                )
            # Todos los audios, incluidos los fragmentos del dictado global,
            # se normalizan a WAV mono de 16 kHz. Así el micrófono puede capturar
            # a 44,1/48 kHz sin escribir una cabecera incorrecta ni romper Whisper.
            prepared = prepare_audio(source, temp_dir)
            result = model_service.transcribe(
                prepared, model_name, language, device, mode
            )
        payload = result.__dict__
        payload["total_seconds"] = round(time.perf_counter() - started, 3)
        payload["request_id"] = g.request_id
        # Guardar en historial
        if payload.get("text"):
            try:
                hist_item = {
                    "id": uuid.uuid4().hex[:12],
                    "text": payload["text"],
                    "model": model_name,
                    "language": payload.get("language", ""),
                    "device": device,
                    "mode": mode,
                    "duration_seconds": payload["total_seconds"],
                    "created_at": time.time(),
                    "created_at_iso": datetime.now().isoformat(timespec="seconds"),
                    "char_count": len(payload["text"]),
                }
                (HISTORY_DIR / f"{hist_item['id']}.json").write_text(
                    json.dumps(hist_item, ensure_ascii=False, indent=2),
                    encoding="utf-8",
                )
                # Limpiar antiguas si excede el máximo
                old = sorted(
                    HISTORY_DIR.glob("*.json"),
                    key=lambda p: p.stat().st_mtime,
                    reverse=True,
                )
                for p in old[HISTORY_MAX:]:
                    p.unlink(missing_ok=True)
            except Exception:
                LOGGER.exception("No se pudo guardar en el historial")
        return jsonify(payload)
    except RuntimeError as exc:
        LOGGER.exception("Error controlado durante transcripción | id=%s", g.request_id)
        message = str(exc)
        code = (
            "DEPENDENCY_ERROR"
            if "dependenc" in message.lower() or "ffmpeg" in message.lower()
            else "TRANSCRIPTION_FAILED"
        )
        return error_response(code, message, 500)
    except Exception as exc:
        LOGGER.exception("Error inesperado durante transcripción | id=%s", g.request_id)
        return error_response(
            "TRANSCRIPTION_FAILED",
            "No se pudo completar la transcripción. Revisa el centro de diagnóstico.",
            500,
            str(exc),
        )


@app.get("/api/diagnostics")
def api_diagnostics():
    try:
        return jsonify(
            {"report": run_diagnostics(include_audio=True), "request_id": g.request_id}
        )
    except Exception as exc:
        LOGGER.exception("Diagnóstico falló")
        return error_response(
            "DIAGNOSTICS_FAILED", "No se pudo ejecutar el diagnóstico.", 500, str(exc)
        )


@app.get("/api/logs")
def api_logs():
    name = request.args.get("name", "app")
    if name not in {"app", "agent", "frontend", "install", "launcher", "tests"}:
        return error_response(
            "LOG_INVALID", "El archivo de log solicitado no existe.", 400
        )
    try:
        limit = int(request.args.get("limit", "300"))
    except ValueError:
        limit = 300
    path = LOG_DIR / f"{name}.log"
    return jsonify(
        {"name": name, "lines": tail_log(path, limit), "request_id": g.request_id}
    )


@app.post("/api/client-log")
def api_client_log():
    data = request.get_json(silent=True) or {}
    level = str(data.get("level", "error")).lower()
    message = str(data.get("message", ""))[:4000]
    context = data.get("context")
    frontend_logger = get_file_logger("frontend")
    log_method = getattr(
        frontend_logger,
        level if level in {"debug", "info", "warning", "error"} else "error",
    )
    log_method(
        "Browser | message=%s | context=%s",
        message,
        json.dumps(context, ensure_ascii=False, default=str)[:6000],
    )
    return jsonify({"ok": True, "request_id": g.request_id})


@app.get("/api/support-bundle")
def api_support_bundle():
    try:
        bundle = create_support_bundle()
        return send_file(
            bundle,
            as_attachment=True,
            download_name=bundle.name,
            mimetype="application/zip",
        )
    except Exception as exc:
        LOGGER.exception("No se pudo crear el paquete de soporte")
        return error_response(
            "SUPPORT_BUNDLE_FAILED",
            "No se pudo crear el paquete de soporte.",
            500,
            str(exc),
        )


@app.post("/api/export/pdf")
def api_export_pdf():
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()
    if not text:
        return error_response("EXPORT_EMPTY", "No hay texto para exportar.", 400)
    filename = secure_filename(data.get("filename", "transcripcion.pdf"))
    if not filename.endswith(".pdf"):
        filename += ".pdf"
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 12, "Feria Transcriber", align="C")
        pdf.ln(18)
        pdf.set_font("Helvetica", "", 11)
        for line in text.split("\n"):
            pdf.multi_cell(0, 6, line)
        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)
        return send_file(
            buf, as_attachment=True, download_name=filename, mimetype="application/pdf"
        )
    except Exception as exc:
        LOGGER.exception("Export PDF fallo")
        return error_response(
            "EXPORT_PDF_FAILED", "No se pudo generar el PDF.", 500, str(exc)
        )


@app.post("/api/export/docx")
def api_export_docx():
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()
    if not text:
        return error_response("EXPORT_EMPTY", "No hay texto para exportar.", 400)
    filename = secure_filename(data.get("filename", "transcripcion.docx"))
    if not filename.endswith(".docx"):
        filename += ".docx"
    try:
        from docx import Document
        from docx.shared import Pt, Inches

        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)
        p = doc.add_paragraph()
        run = p.add_run("Feria Transcriber")
        run.bold = True
        run.font.size = Pt(18)
        p.alignment = 1
        doc.add_paragraph()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as exc:
        LOGGER.exception("Export DOCX fallo")
        return error_response(
            "EXPORT_DOCX_FAILED", "No se pudo generar el documento Word.", 500, str(exc)
        )


@app.get("/api/history")
def api_history():
    """Lista el historial de transcripciones."""
    items = []
    try:
        for f in sorted(
            HISTORY_DIR.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True
        )[:HISTORY_MAX]:
            try:
                data = json.loads(f.read_text(encoding="utf-8"))
                items.append(data)
            except Exception:
                continue
    except Exception as exc:
        LOGGER.exception("No se pudo leer el historial")
        return error_response(
            "HISTORY_READ_FAILED", "No se pudo leer el historial.", 500, str(exc)
        )
    return jsonify({"items": items, "request_id": g.request_id})


@app.post("/api/history")
def api_history_save():
    """Guarda una transcripción en el historial."""
    data = request.get_json(silent=True) or {}
    text = (data.get("text") or "").strip()
    if not text:
        return error_response("HISTORY_EMPTY", "El texto está vacío.", 400)
    item = {
        "id": uuid.uuid4().hex[:12],
        "text": text,
        "model": data.get("model", "unknown"),
        "language": data.get("language", ""),
        "device": data.get("device", "auto"),
        "mode": data.get("mode", "manual"),
        "duration_seconds": data.get("duration_seconds", 0),
        "created_at": time.time(),
        "created_at_iso": datetime.now().isoformat(timespec="seconds"),
        "char_count": len(text),
    }
    try:
        path = HISTORY_DIR / f"{item['id']}.json"
        path.write_text(
            json.dumps(item, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        return jsonify({"item": item, "request_id": g.request_id})
    except Exception as exc:
        LOGGER.exception("No se pudo guardar en el historial")
        return error_response(
            "HISTORY_SAVE_FAILED", "No se pudo guardar.", 500, str(exc)
        )


@app.delete("/api/history/<item_id>")
def api_history_delete(item_id: str):
    """Elimina una entrada del historial."""
    path = HISTORY_DIR / f"{item_id}.json"
    if path.exists():
        path.unlink()
        return jsonify({"deleted": True, "request_id": g.request_id})
    return error_response("HISTORY_NOT_FOUND", "Entrada no encontrada.", 404)


@app.get("/api/gpu-stats")
def api_gpu_stats():
    """Estadísticas de la GPU (VRAM, temperatura, uso)."""
    stats = {"cuda_available": False, "devices": []}
    try:
        import ctranslate2

        count = ctranslate2.get_cuda_device_count()
        stats["cuda_available"] = count > 0
        if count > 0:
            for i in range(count):
                info = {"index": i, "name": f"GPU {i}"}
                try:
                    out = subprocess.run(
                        [
                            "nvidia-smi",
                            f"--id={i}",
                            "--query-gpu=name,memory.used,memory.total,utilization.gpu,temperature.gpu",
                            "--format=csv,noheader,nounits",
                        ],
                        capture_output=True,
                        text=True,
                        timeout=5,
                    )
                    if out.returncode == 0 and out.stdout.strip():
                        parts = [p.strip() for p in out.stdout.strip().split(",")]
                        if len(parts) >= 5:
                            info["name"] = parts[0]
                            info["memory_used_mb"] = int(float(parts[1]))
                            info["memory_total_mb"] = int(float(parts[2]))
                            info["gpu_util_pct"] = int(float(parts[3]))
                            info["temperature_c"] = int(float(parts[4]))
                except (FileNotFoundError, subprocess.TimeoutExpired, Exception):
                    pass
                stats["devices"].append(info)
    except Exception:
        pass
    return jsonify({"stats": stats, "request_id": g.request_id})


@app.post("/api/notify")
def api_notify():
    """Envía una notificación al sistema (solo Windows)."""
    data = request.get_json(silent=True) or {}
    title = str(data.get("title", "Feria Transcriber"))[:100]
    message = str(data.get("message", ""))[:300]
    if os.name == "nt":
        try:
            from windows_toasts import WindowsToaster, Toast

            toaster = WindowsToaster(title)
            toast = Toast()
            toast.text_fields = [message]
            toaster.show_toast(toast)
        except ImportError:
            try:
                import ctypes

                ctypes.windll.user32.MessageBeep(0x40)
            except Exception:
                pass
        except Exception:
            pass
    return jsonify({"ok": True, "request_id": g.request_id})


@app.get("/api/languages")
def api_languages():
    """Lista de idiomas soportados por Whisper."""
    return jsonify(
        {"languages": LANGUAGE_NAMES, "default": "es", "request_id": g.request_id}
    )


@app.post("/api/open-logs")
def api_open_logs():
    try:
        if os.name == "nt":
            os.startfile(LOG_DIR)
        else:
            subprocess.Popen(["xdg-open", str(LOG_DIR)])
        return jsonify({"ok": True, "request_id": g.request_id})
    except Exception as exc:
        LOGGER.exception("No se pudo abrir la carpeta de logs")
        return error_response(
            "OPEN_LOGS_FAILED", "No se pudo abrir la carpeta de logs.", 500, str(exc)
        )


@app.errorhandler(RequestEntityTooLarge)
def file_too_large(_error):
    return error_response(
        "FILE_TOO_LARGE", "El archivo supera el límite de 1,5 GB.", 413
    )


@app.errorhandler(404)
def not_found(_error):
    if request.path.startswith("/api/"):
        return error_response("NOT_FOUND", "La ruta solicitada no existe.", 404)
    return "No encontrado", 404


@app.errorhandler(Exception)
def unexpected_error(exc):
    LOGGER.exception("Error global no controlado | id=%s", getattr(g, "request_id", ""))
    return error_response(
        "INTERNAL_ERROR",
        "Se produjo un error interno. El detalle quedó guardado en logs/app.log.",
        500,
        str(exc),
    )


def find_available_port(preferred: int = 5000) -> int:
    for port in range(preferred, preferred + 30):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
            try:
                sock.bind(("127.0.0.1", port))
                return port
            except OSError:
                continue
    raise RuntimeError("No hay un puerto local disponible entre 5000 y 5029.")


def open_browser(url: str) -> None:
    try:
        webbrowser.open(url, new=2)
    except Exception:
        LOGGER.exception("No se pudo abrir el navegador")


def main() -> int:
    try:
        from waitress import serve
    except Exception:
        LOGGER.exception("Waitress no está instalado")
        print("ERROR: faltan dependencias. Ejecuta instalar.bat.")
        return 1
    port = find_available_port(5000)
    url = f"http://127.0.0.1:{port}"
    agent_controller.configure(url)
    agent_controller.start_monitor()
    config = load_config()
    (RUNTIME_DIR / "server.json").write_text(
        json.dumps(
            {"url": url, "port": port, "pid": os.getpid(), "started_at": time.time()},
            indent=2,
        ),
        encoding="utf-8",
    )
    LOGGER.info("Servidor iniciando | url=%s", url)
    print("\n" + "=" * 68)
    print("  FERIA TRANSCRIBER 2.2")
    print(f"  Web: {url}")
    print(f"  Logs: {LOG_DIR}")
    print(
        f"  Dictado global: {'activado' if config['dictation']['enabled'] else 'desactivado'} · {config['dictation']['hotkey'].upper()}"
    )
    print("  Cierra esta ventana para detener todo.")
    print("=" * 68 + "\n")
    if config["web"]["open_browser"]:
        threading.Timer(1.0, open_browser, args=(url,)).start()
    if config["dictation"]["enabled"]:
        threading.Timer(1.4, agent_controller.start).start()
    threading.Timer(
        1.8, model_service.warmup_async, args=(config["model"], config["device"])
    ).start()
    try:
        serve(app, host="127.0.0.1", port=port, threads=8, channel_timeout=3600)
        return 0
    except KeyboardInterrupt:
        return 0
    finally:
        agent_controller.stop()


if __name__ == "__main__":
    raise SystemExit(main())
